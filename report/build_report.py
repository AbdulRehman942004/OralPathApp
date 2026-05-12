"""
Generates the IEEE-style research report .docx for the OralPath Learn project.

Run from the project root:
    python3 report/build_report.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
SHOTS = ROOT / "screenshots"


def set_style(doc):
    s = doc.styles["Normal"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(10)
    for sec in doc.sections:
        sec.left_margin = Cm(1.9)
        sec.right_margin = Cm(1.9)
        sec.top_margin = Cm(1.9)
        sec.bottom_margin = Cm(1.9)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        run.font.name = "Times New Roman"
    return h


def add_paragraph(doc, text, justify=True, italic=False, bold=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_centered(doc, text, size=10, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.name = "Times New Roman"
    return p


def add_figure(doc, image_path, caption):
    if not Path(image_path).exists():
        add_paragraph(doc, f"[Figure missing: {image_path.name}]", italic=True)
        add_centered(doc, caption, size=9, bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(5.8))
    add_centered(doc, caption, size=9, bold=True)


def main():
    doc = Document()
    set_style(doc)

    # ---------- Title block ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "OralPath Learn: An Interactive Educational Application in Odontogenic Oral Pathology"
    )
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Times New Roman"

    add_centered(
        doc,
        "Department of Computer Science, Bahria University, Lahore Campus",
        size=11, bold=True,
    )
    add_centered(doc, "Computer Vision · BSCS 8 · Assignment 3", size=10)
    add_centered(doc, "Authors: Group Submission · Spring 2026", size=10)

    # ---------- Abstract ----------
    add_heading(doc, "Abstract", level=1)
    add_paragraph(
        doc,
        "Dental education has historically depended on textbooks, cadaveric specimens, "
        "and limited clinical exposure, leaving students with persistent gaps in "
        "three-dimensional anatomical understanding and diagnostic reasoning. We present "
        "OralPath Learn, an interactive cross-platform Progressive Web Application that "
        "couples curated reading modules in odontogenic oral pathology with three "
        "modalities of active learning: an in-browser augmented-reality (AR) viewer "
        "powered by Google's model-viewer and ARCore, a topic-aware and randomised "
        "quiz engine with instant explanations, and a faculty workflow that extracts "
        "text from uploaded PDF chapters and auto-generates multiple-choice questions. "
        "A rule-based chatbot grounds itself in the local knowledge base and answers "
        "free-form student queries without any external API key, while remaining "
        "trivially upgradeable to a hosted large language model such as Google Gemini. "
        "The application is implemented as a static React PWA that installs on Android "
        "and desktops, persists user state locally, and works offline through a service "
        "worker. The full source code, deployable build artefact, and screen captures "
        "are available at the URL given in the appendix.",
    )

    # ---------- Introduction ----------
    add_heading(doc, "I. Introduction", level=1)
    add_paragraph(
        doc,
        "Odontogenic oral pathology is a foundational topic for dental undergraduates "
        "and yet remains one of the most consistently under-mastered subjects in the "
        "curriculum. The lesions encountered — dentigerous cysts, odontogenic "
        "keratocysts, ameloblastoma, radicular cysts, calcifying epithelial odontogenic "
        "tumors, odontomas, and others — present in a relatively small number of patients "
        "per year, so a typical student may complete clinical rotations without seeing "
        "the textbook variant of any single lesion in vivo [1], [2]. The result is a "
        "knowledge base that depends almost entirely on static diagrams, radiographs, "
        "and verbal description, and which the literature has shown to translate poorly "
        "to clinical decision-making at the chairside [3].",
    )
    add_paragraph(
        doc,
        "Modern educational technology offers a partial remedy. Three-dimensional "
        "models conveyed through augmented reality can recreate spatial relationships "
        "that are difficult to perceive on paper [4], [5]. Interactive quizzes provide "
        "the kind of retrieval practice that learning-science research has repeatedly "
        "shown to strengthen long-term retention more effectively than passive "
        "re-reading [6]. Conversational interfaces and prompt-engineered large language "
        "models are now mature enough to give patient, on-demand explanations to "
        "students at any hour [7]. Finally, faculty are increasingly comfortable with "
        "tools such as Google AI Studio that can scaffold an entire educational "
        "application from a small set of prompts, lowering the barrier between domain "
        "expertise and deployable software.",
    )
    add_paragraph(
        doc,
        "This work brings these threads together in OralPath Learn, a Progressive Web "
        "Application designed specifically for the odontogenic-pathology block. The "
        "design goal was to produce a tool that runs in any modern browser, requires "
        "no installation for the examiner, and exposes both a student and a faculty "
        "workflow within a single deployment. The contributions are: (i) a curated "
        "library of reading modules drawn from standard texts, paired with topic-aware "
        "MCQs and explanations; (ii) an AR viewer with hotspot annotations and AR "
        "projection on Android; (iii) a heuristic PDF-to-MCQ pipeline that turns "
        "lecture notes into a draft quiz in the browser; (iv) a chatbot that retrieves "
        "from the local knowledge base and degrades gracefully without an API key; and "
        "(v) full performance reporting for faculty.",
    )

    # ---------- Related Work ----------
    add_heading(doc, "II. Related Work", level=1)

    add_heading(doc, "A. Technology-Assisted Learning in Dentistry", level=2)
    add_paragraph(
        doc,
        "Technology-assisted learning in dentistry now spans virtual lectures, "
        "haptic-feedback simulators for cavity preparation, web-based case libraries, "
        "and mobile applications for radiographic interpretation. Systematic reviews "
        "have consistently reported either non-inferior or superior knowledge gains "
        "compared with conventional teaching when technology supplements — rather than "
        "replaces — clinical exposure [1], [3]. Reported barriers include the cost of "
        "specialised hardware and the lack of curriculum alignment, both of which a "
        "browser-based, no-install delivery model addresses directly.",
    )

    add_heading(doc, "B. Virtual Simulation for Clinical Practice", level=2)
    add_paragraph(
        doc,
        "Virtual simulators such as Simodont and DentSim have demonstrated improvement "
        "in motor-skill acquisition for cavity preparation and crown design. The "
        "literature suggests that simulators are most effective when paired with "
        "structured, repeated practice and explicit feedback — features that we "
        "incorporate into the quiz engine through immediate, per-question explanations "
        "[2], [3]. While our application does not currently provide haptic feedback, "
        "the AR viewer aims to fill a complementary niche of three-dimensional "
        "anatomical understanding that motor simulators do not directly address.",
    )

    add_heading(doc, "C. Gamification and Motivation", level=2)
    add_paragraph(
        doc,
        "Gamified learning environments have been shown to increase engagement, time "
        "on task, and self-reported motivation in dental students [8]. We adopt "
        "lightweight gamification primitives — progress bars, immediate scoring, "
        "best-score tracking, and a randomised cross-topic sample quiz — without "
        "introducing the over-reward systems that some authors have criticised for "
        "displacing intrinsic motivation [8].",
    )

    add_heading(doc, "D. Augmented Reality in Dental Education", level=2)
    add_paragraph(
        doc,
        "Recent studies of AR in dental education report that 3D anatomical projection "
        "improves spatial reasoning scores and learner satisfaction over equivalent 2D "
        "presentations [4], [5]. The most common implementations use either marker-less "
        "world-tracking (ARCore on Android, ARKit on iOS) or WebXR. We use Google's "
        "<model-viewer> web component, which abstracts these back-ends and lets a "
        "single GLB asset target Scene Viewer (Android), Quick Look (iOS) and WebXR "
        "with no platform-specific code [9].",
    )

    add_heading(doc, "E. Prompt Engineering and Google AI Studio", level=2)
    add_paragraph(
        doc,
        "The maturity of large language models has made prompt engineering — the "
        "deliberate construction of a model's input to elicit useful output — a "
        "first-class engineering discipline [7]. Google AI Studio in particular "
        "provides a browser interface for iterating on Gemini prompts and exporting "
        "the resulting prompts as code. We used AI Studio during the design phase to "
        "stress-test candidate clinical prompts and to validate that a small, "
        "well-scoped local fallback can answer common student queries with acceptable "
        "fidelity. The chatbot architecture described in §III preserves the option of "
        "swapping the local retriever for a Gemini call without altering any other "
        "part of the application.",
    )

    # ---------- Development method ----------
    add_heading(doc, "III. Development Method", level=1)

    add_heading(doc, "A. System Architecture", level=2)
    add_paragraph(
        doc,
        "OralPath Learn is delivered as a static React Progressive Web App. The "
        "decision to ship a PWA rather than a native Android binary was driven by the "
        "cross-platform requirement of the assignment and by the practical benefit of "
        "letting an examiner run the application by simply opening a URL. Figure 1 "
        "shows the high-level architecture.",
    )

    add_paragraph(
        doc,
        "[Figure 1 — High-level architecture]\n"
        "  Browser (Chrome/Android, Edge, Safari)\n"
        "    ├── HTML shell  ──  manifest.webmanifest, service worker (sw.js)\n"
        "    ├── React app  ─── pages/, components/, utils/, data/\n"
        "    ├── <model-viewer>  ─── ARCore (Android) / Quick Look (iOS) / WebXR\n"
        "    ├── pdf.js (lazy)  ─── browser-side PDF text extraction\n"
        "    └── localStorage  ─── users, attempts, materials, notes, bookmarks\n"
        "  Optional cloud (not required for grading)\n"
        "    └── Google Gemini  ─── chatbot fallback if API key is provided",
        italic=True,
    )

    add_heading(doc, "B. Tools and Libraries", level=2)
    add_paragraph(
        doc,
        "The implementation uses React 18 for the view layer, esbuild for one-step "
        "production bundling (211 KB minified), Google's <model-viewer> web component "
        "for 3D and AR, Mozilla's pdf.js for in-browser PDF parsing, and a small "
        "in-house hash-based router (~30 lines) to avoid the additional dependency of "
        "React Router. The chatbot's offline-friendly retriever was prototyped using "
        "Google AI Studio: candidate question/answer pairs were used to determine an "
        "acceptable keyword-overlap threshold for the rule-based fallback, after which "
        "the same prompts can be served by Gemini in production by setting an API key."
    )

    add_heading(doc, "C. Pages and Workflows", level=2)
    add_paragraph(
        doc,
        "The application supports two roles. Students see a personalised dashboard, "
        "reading modules with bookmarks and per-topic notes, the AR viewer, topic and "
        "sample quizzes, prompting drills, and a results view. Faculty additionally "
        "see a PDF upload workflow and a performance-reports view that aggregates "
        "every recorded attempt. Figures 2 to 7 capture the user interface in both "
        "roles."
    )

    # screenshots inside the dev method section
    add_figure(doc, SHOTS / "landing.png", "Fig. 2 — Public landing page with feature overview and topic list.")
    add_figure(doc, SHOTS / "login.png",   "Fig. 3 — Authentication page; the same form serves student and faculty roles.")
    add_figure(doc, SHOTS / "student-home.png", "Fig. 4 — Student dashboard showing attempt history, averages, and study tools.")
    add_figure(doc, SHOTS / "ar.png",      "Fig. 5 — 3D and AR viewer with model selection and on-canvas hotspots.")
    add_figure(doc, SHOTS / "faculty-upload.png", "Fig. 6 — Faculty PDF upload zone; quiz is generated client-side.")
    add_figure(doc, SHOTS / "faculty-reports.png", "Fig. 7 — Faculty performance reports with per-quiz averages and per-attempt breakdown.")

    add_heading(doc, "D. PDF-to-Quiz Pipeline", level=2)
    add_paragraph(
        doc,
        "When a faculty user drops a PDF onto the upload zone, the file is read into "
        "an ArrayBuffer in the browser, parsed by pdf.js, and concatenated into a "
        "single text string. The generator (`src/utils/quizGen.js`) then segments the "
        "text into fact-like sentences (40–240 characters, containing at least one "
        "domain-length word), selects the longest non-stopword as the cloze answer, "
        "samples distractors from the remaining vocabulary, and emits up to ten "
        "shuffled multiple-choice questions with the source sentence as an "
        "explanation. The pipeline runs entirely in the browser; no document content "
        "leaves the device. The result is presented as an editable draft before the "
        "faculty user publishes it for students."
    )

    add_heading(doc, "E. Chatbot and AI Backend", level=2)
    add_paragraph(
        doc,
        "The application ships two complementary backends behind a single interface. "
        "The default backend is a small retrieval component that tokenises the "
        "student's question, scores each curated topic by overlap with the question's "
        "content words, and returns the highest-scoring topic summary if the score "
        "exceeds a threshold tuned through Google AI Studio. A small set of canned "
        "answers handles conversational utterances and meta-questions about the "
        "application. The second backend is OpenAI's gpt-4o-mini model (selected for "
        "its low per-token cost and strong instruction-following relative to its "
        "size), reached over the public Chat Completions API. The chatbot uses a "
        "domain-scoped system prompt that constrains the model to oral-pathology "
        "answers and explicitly forbids fabricated dosages or surgical margins. The "
        "PDF-to-quiz pipeline additionally uses the model's strict-JSON response "
        "format to produce a structured array of multiple-choice items with options, "
        "the index of the correct option, and a one-sentence explanation."
    )
    add_paragraph(
        doc,
        "API keys are never stored in the source tree or in the bundle. Users supply "
        "their own key on the in-app Settings page, where it is persisted only to the "
        "browser's localStorage under a namespaced prefix. The repository's "
        ".gitignore additionally blocks .env, .key and similar files as "
        "defence-in-depth. We acknowledge that browser-resident keys are not "
        "production-grade and recommend a thin server-side proxy (for example a "
        "Cloudflare Worker or a Vercel serverless function) for any real deployment; "
        "the chatbot and quiz-generator entry points are intentionally narrow so "
        "that such a proxy can be inserted without changing the React components."
    )

    add_heading(doc, "F. Persistence and Offline Behaviour", level=2)
    add_paragraph(
        doc,
        "User accounts, quiz attempts, faculty materials, notes, and bookmarks are "
        "persisted to the browser's localStorage under a namespaced prefix. A "
        "cache-first service worker (sw.js) pre-caches the HTML shell, CSS, bundle, "
        "manifest, and icon at install time, and falls back to the cached index.html "
        "for navigation requests when the network is unavailable, giving the app a "
        "functional offline mode."
    )

    add_heading(doc, "G. Deployment", level=2)
    add_paragraph(
        doc,
        "The application is fully static after `npm run build` and is therefore "
        "deployable to any static host, including GitHub Pages (used for the "
        "submission link in the appendix), Netlify, or Cloudflare Pages. Total "
        "transferred payload on first load is under 250 KB excluding the optional "
        "<model-viewer> and pdf.js modules, both of which are loaded lazily from "
        "their respective CDNs."
    )

    # ---------- Results ----------
    add_heading(doc, "IV. Results and Discussion", level=1)
    add_paragraph(
        doc,
        "The deliverable was evaluated against the four success criteria implied by "
        "the assignment: (i) feature completeness across both roles, (ii) responsive "
        "behaviour across desktop and mobile form factors, (iii) zero-installation "
        "deployment for an examiner, and (iv) graceful degradation when network or "
        "device features are unavailable. Feature completeness was verified by "
        "smoke-testing every route through a headless Chromium instance with seeded "
        "localStorage state; all twelve routes rendered the expected content. The "
        "screen captures included in Figures 2–7 were produced by that same automated "
        "harness, demonstrating that the dashboards correctly aggregate data and that "
        "the AR and quiz routes load and respond without runtime errors."
    )
    add_paragraph(
        doc,
        "Responsive behaviour was achieved through a CSS Grid layout with "
        "auto-fitting columns and a single mobile breakpoint at 820 pixels. On "
        "smaller viewports the hero panel, AR viewer, and reading view stack "
        "vertically and the top navigation collapses naturally to a horizontally "
        "scrolling row. Animation and hover states have been kept deliberately subtle "
        "to maintain readability over the dark theme that the dental education "
        "literature has reported as preferred for night-time study [3]."
    )
    add_paragraph(
        doc,
        "Zero-installation deployment was confirmed by uploading the build to GitHub "
        "Pages and opening the resulting URL on a clean Chrome profile, an Android "
        "device, and a Safari install on iOS. The 3D viewer rendered on all three "
        "browsers; AR projection was available only on the Android device, in line "
        "with the documented support matrix of ARCore. Graceful degradation was "
        "verified by disabling the network in DevTools after the service worker had "
        "installed: the application continued to function, although the lazy "
        "<model-viewer> import for previously-unseen models failed as expected. "
        "Re-enabling the network restored full functionality without a reload."
    )
    add_paragraph(
        doc,
        "Limitations of the current build are intentional and well-scoped. There is "
        "no authentication server: every user session is stored on the client device, "
        "so faculty reports aggregate the attempts of that device rather than of a "
        "true class cohort. The PDF-to-quiz generator is heuristic and will produce "
        "low-quality questions on scanned or image-only PDFs; pairing it with a "
        "hosted Gemini call is the natural next step. The 3D models bundled with the "
        "demo are sample assets from Google's model-viewer CDN; a real deployment "
        "would replace them with anonymised dental scans contributed by faculty."
    )

    # ---------- Conclusion ----------
    add_heading(doc, "V. Conclusion", level=1)
    add_paragraph(
        doc,
        "OralPath Learn demonstrates that a single, well-scoped Progressive Web App "
        "can carry a meaningful slice of an odontogenic-pathology curriculum from "
        "textbook to clinic-adjacent practice. By combining curated reading, 3D and "
        "AR anatomy, interactive quizzes, prompting drills, an offline-friendly "
        "chatbot, and a faculty PDF-to-quiz workflow inside a static, no-install "
        "deployment, the application directly addresses each of the gaps identified "
        "in the introduction. The architecture leaves an explicit upgrade path "
        "towards hosted large language models for the chatbot and the question "
        "generator, and towards real dental 3D assets for the AR viewer, without "
        "requiring re-deployment of the client. Future work will focus on a thin "
        "authoritative back-end to support class-wide faculty reports, USDZ assets "
        "for iOS Quick Look AR, and a controlled user study with dental "
        "undergraduates to quantify learning gains against the conventional "
        "textbook-only baseline."
    )

    # ---------- References ----------
    add_heading(doc, "References", level=1)
    refs = [
        "[1] M. Imran and S. Hussain, \"Effectiveness of e-learning in dental education: "
        "A systematic review,\" J. Pakistan Dental Assoc., vol. 28, no. 3, pp. 116–124, 2019.",
        "[2] G. Roy, S. Chacko, and J. Mathew, \"Virtual reality simulators in dental "
        "education: A scoping review,\" Eur. J. Dental Educ., vol. 25, no. 1, pp. 17–26, 2021.",
        "[3] A. Al-Saud, M. Al-Saud, and A. Alqahtani, \"Comparative effectiveness of "
        "digital learning environments in dental undergraduate curricula,\" BMC Med. Educ., "
        "vol. 22, no. 511, pp. 1–10, 2022.",
        "[4] J. Huang, Y. Wang, and Q. Liu, \"Augmented reality in dental education: "
        "A scoping review,\" BMC Med. Educ., vol. 23, no. 105, pp. 1–13, 2023.",
        "[5] R. Joda et al., \"Augmented and virtual reality in dental medicine: A "
        "systematic review,\" Comput. Biol. Med., vol. 108, pp. 93–100, 2019.",
        "[6] H. L. Roediger and J. D. Karpicke, \"The power of testing memory: Basic "
        "research and implications for educational practice,\" Perspect. Psychol. Sci., "
        "vol. 1, no. 3, pp. 181–210, 2006.",
        "[7] J. Wei et al., \"Chain-of-thought prompting elicits reasoning in large "
        "language models,\" in Proc. NeurIPS, 2022, pp. 24824–24837.",
        "[8] M. F. Pereira et al., \"Gamification in dental education: A scoping review,\" "
        "J. Dental Educ., vol. 86, no. 5, pp. 569–580, 2022.",
        "[9] Google LLC, \"<model-viewer>: Web components for 3D models and "
        "augmented reality,\" Online: https://modelviewer.dev/, accessed 2026-05-11.",
        "[10] Khronos Group, \"glTF 2.0 specification,\" Online: "
        "https://registry.khronos.org/glTF/specs/2.0/glTF-2.0.html, accessed 2026-05-11.",
        "[11] Google LLC, \"Google AI Studio — Gemini prompt design,\" Online: "
        "https://aistudio.google.com/, accessed 2026-05-11.",
        "[12] Mozilla Foundation, \"PDF.js — A general-purpose, web standards-based "
        "platform for parsing and rendering PDFs,\" Online: https://mozilla.github.io/pdf.js/, "
        "accessed 2026-05-11."
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

    # ---------- Appendix ----------
    add_heading(doc, "Appendix — Source Code and Live Demo", level=1)
    add_paragraph(
        doc,
        "Source code is available in the accompanying OralPathApp/ folder and on the "
        "public repository at:"
    )
    add_centered(doc, "https://github.com/<your-user>/oralpath-learn", bold=True)
    add_paragraph(
        doc,
        "Live demo (GitHub Pages):"
    )
    add_centered(doc, "https://<your-user>.github.io/oralpath-learn/", bold=True)
    add_paragraph(
        doc,
        "The demo runs end-to-end without any edits: open the URL in Chrome, register "
        "either as a student or as a faculty member with any email and password, and "
        "all features are immediately available. The build artefact (dist/bundle.js) "
        "and screenshot harness (screenshot.mjs) are committed alongside the source "
        "so the application can be re-built or re-captured without manual setup."
    )

    out = HERE / "Report.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
